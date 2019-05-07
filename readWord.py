#!/usr/bin/env python
# -*- coding:utf-8 -*-

'''
@File  : readWord.py
@Author: Grace
@Date  : 2019/4/26
@Desc  : 
'''

# import modules
# import docx
from docx import Document
import sys
import os

# global variables
path = '/Users/yongqi/PycharmProjects/RiskAnalysis/wenshu/'

keywords =['驳回上诉','维持原判','维持原裁定']

# class definition

# function definition
def ReadDocx():
    reload(sys)
    sys.setdefaultencoding('utf-8')

    # 存在质量风险的案件总数量
    quality_counter = 0
    # 管辖权
    guanxia_counter = 0
    # 不予受理
    shouli_counter = 0
    # 股权 或 股份
    guquan_counter = 0
    # 借贷 或 借款
    jiedai_counter = 0
    # 合同
    hetong_counter = 0
    # 执行人执行异议
    zhixing_counter = 0
    # 赔偿相关
    peichang_counter = 0
    # 不正当竞争
    compete_counter = 0
    # 交通事故/医疗费用
    shigu_counter = 0
    # 债权
    zhaiquan_counter = 0
    # 房屋买卖
    fangwu_counter = 0
    # 著作权/专利权
    zhuzuo_counter = 0

    ## 其他风险


    for filename in os.listdir(path):
        if os.path.splitext(filename)[1] == '.docx':
            # print '1'
            try:

                document = Document(path+filename)
            except Exception:
                print 'doc 文档.....'
                print filename
                continue
            else:
                # # 读取文档中所有的段落列表
                # ps = document.paragraphs
                # # 每个段落有两个属性：style和text
                # ps_content = [(x.text, x.style.name) for x in ps]
                # with open('out.tmp', 'w+') as fout:
                #     fout.write('')
                # # 读取段落并写入一个文件
                # with open('out.tmp', 'a+') as fout:
                #     for p in ps_content:
                #         fout.write(p[0] + '\t' + p[1] + '\n\n')
                #
                # # 读取文档中的所有段落的列表
                # tables = document.tables
                # # 遍历table，并将所有单元格内容写入文件中
                # with open('out.tmp', 'a+') as fout:
                #     for table in tables:
                #         for row in table.rows:
                #             for cell in row.cells:
                #                 fout.write(cell.text + '\t')
                #             fout.write('\n')


                # document = Document(path + '23.docx')
                # 读取文档中所有的段落列表
                ps = document.paragraphs
                # 每个段落有两个属性：style和text
                ps_content = [(x.text, x.style.name) for x in ps]

                flag = 1 # 存在质量风险
                for p in ps_content:

                    content = str(p[0] + '\t' + p[1])
                    if any(keyword in content for keyword in keywords) == True:
                        flag = 0 # 不存在质量风险
                        break
                if flag == 1:
                    quality_counter = quality_counter + 1

                # 存在质量风险，检查是否存在"管辖权"关键词
                if flag == 1:
                    for p in ps_content:
                        content = str(p[0] + '\t' + p[1])
                        stop = 0
                        if ('管辖权' in content) == True:
                            guanxia_counter = guanxia_counter + 1
                            stop = 1
                        if ('不予受理' in content) == True:
                            shouli_counter = shouli_counter + 1
                            stop = 1
                        if any(w in content for w in ['股权','股份']) == True:
                            guquan_counter = guquan_counter + 1
                            stop = 1
                        if any(w in content for w in ['借贷', '借款']) == True:
                            jiedai_counter = jiedai_counter + 1
                            stop = 1
                        if ('合同' in content) == True:
                            hetong_counter = hetong_counter + 1
                            stop = 1
                        if ('执行异议' in content) == True:
                            zhixing_counter = zhixing_counter + 1
                            stop = 1
                        if ('赔偿' in content) == True:
                            peichang_counter = peichang_counter + 1
                            stop = 1
                        if ('竞争' in content) == True:
                            compete_counter = compete_counter + 1
                            stop = 1
                        if any(w in content for w in ['交通事故', '医疗费用']) == True:
                            shigu_counter = shigu_counter + 1
                            stop = 1
                        if ('债权' in content) == True:
                            zhaiquan_counter = zhaiquan_counter + 1
                            stop = 1
                        if ('房屋' in content) == True:
                            fangwu_counter = fangwu_counter + 1
                            stop = 1
                        if any(w in content for w in ['著作权', '专利权']) == True:
                            zhuzuo_counter = zhuzuo_counter + 1
                            stop = 1
                        if stop == 1:
                            break
        elif os.path.splitext(filename)[1] == '.txt':
            with open(path+filename) as fn:
                text = fn.read()
                flag = 1
                if any(keyword in text for keyword in keywords) == True:
                    flag = 0 # 不存在质量风险
                if flag == 1:
                    quality_counter = quality_counter + 1
                if ('管辖权' in text) == True:
                    guanxia_counter = guanxia_counter + 1
                if ('不予受理' in text) == True:
                    shouli_counter = shouli_counter + 1
                if any(w in text for w in ['股权','股份']) == True:
                    guquan_counter = guquan_counter + 1
                if any(w in text for w in ['借贷', '借款']) == True:
                    jiedai_counter = jiedai_counter + 1
                if ('合同' in text) == True:
                    hetong_counter = hetong_counter + 1
                if ('执行异议' in text) == True:
                    zhixing_counter = zhixing_counter + 1
                if ('赔偿' in text) == True:
                    peichang_counter = peichang_counter + 1
                if ('竞争' in text) == True:
                    compete_counter = compete_counter + 1
                if any(w in text for w in ['交通事故','医疗费用']) == True:
                    shigu_counter = shigu_counter + 1
                if ('债权' in text) == True:
                    zhaiquan_counter = zhaiquan_counter + 1
                if ('房屋' in text) == True:
                    fangwu_counter = fangwu_counter + 1
                if any(w in text for w in ['著作权','专利权']) == True:
                    zhuzuo_counter = zhuzuo_counter + 1

    print '质量风险案件总数:',quality_counter
    print '涉及到管辖权:{:.3%}'.format(float(guanxia_counter)/float(quality_counter))
    print '存在不予受理情况:{:.3%}'.format(float(shouli_counter)/float(quality_counter))
    print '涉及到股权、股份:{:.3%}'.format(float(guquan_counter)/float(quality_counter))
    print '涉及到借贷、借款:{:.3%}'.format(float(jiedai_counter)/float(quality_counter))
    print '涉及到合同:{:.3%}'.format(float(hetong_counter)/float(quality_counter))
    print '（存疑）对执行人执行异议提出申诉:{:.3%}'.format(float(zhixing_counter)/float(quality_counter))
    print '涉及到赔偿:{:.3%}'.format(float(peichang_counter)/float(quality_counter))
    print '涉及到不正当竞争关系:{:.3%}'.format(float(compete_counter)/float(quality_counter))
    print '涉及到交通事故、医疗费用赔偿:{:.3%}'.format(float(shigu_counter)/float(quality_counter))
    print '涉及到债权（转让等事宜）:{:.3%}'.format(float(zhaiquan_counter)/float(quality_counter))
    print '涉及到房屋:{:.3%}'.format(float(fangwu_counter)/float(quality_counter))
    print '涉及到著作权或专利权保护:{:.3%}'.format(float(zhuzuo_counter)/float(quality_counter))




# main function
if __name__ == '__main__':
    ReadDocx()
